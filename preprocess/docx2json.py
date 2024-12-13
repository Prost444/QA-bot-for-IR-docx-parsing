from docx.table import Table as TableObject
from docx import Document as CreateDocument
from docx.document import Document as DocumentObject
from pathlib import Path
import json
import spacy
from typing import Generator

def iter_docx_data(page_name : str, doc : DocumentObject, nlp : spacy.Language) -> Generator[dict[str,object],None,None]:
	containers = doc.tables.copy()

	containers = []

	for table in doc.tables:
		values = []

		for row in table.rows:
			row_values = []
			
			for cell in row.cells:
				row_values.append(cell.text)

			values.append(row_values)

		containers.append((table._element,values))

	current_list = []
	first_element = None

	for paragraph in doc.paragraphs:
		if "List" in paragraph.style.name:
			if not current_list:
				first_element = paragraph._element

			current_list.append([paragraph.text])
		elif current_list:
			containers.append((first_element,current_list))
			current_list = []
			first_element = None
			
	for i in range(len(containers)):
		container_element,container_values = containers[i][0],containers[i][1]

		full_context = []

		for element in doc.element.body:
			if element == container_element:
				break

			if element.text:
				full_context.append(element.text)

		data = {
			'uuid': f'wiki-{1}-{"_".join(page_name.split())}-{i + 1}',
			"context_before": [],
			"caption": "",
			"header": [],
			"data": [],
		}
	
		data["data"] = container_values[1:]
		data["header"] = container_values[0]

		for paragraph in full_context[-3:]:
			paragraph_sents = []
			parsed_paragraph = nlp(paragraph)
			#print('Paragraph: ', parsed_paragraph)
			for sent in parsed_paragraph.sents:
				paragraph_sents.append(sent.text)
				#print('Sent: ', sent)
			data["context_before"].append(paragraph_sents)

		yield data
def docx_to_json(files : list[str], nlp : spacy.Language) -> str:
	examples = []

	for file in files:
		filepath = Path(file)

		docx = CreateDocument(filepath)
			
		for data in iter_docx_data(filepath.stem, docx, nlp):
			examples.append(data)

	return json.dumps(examples)
