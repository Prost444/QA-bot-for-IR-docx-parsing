import faiss
import os
from docx.table import Table as TableObject
from docx import Document as CreateDocument
from docx.document import Document as DocumentObject
from langchain.docstore.document import Document as LangchainDocument
from langchain_community.vectorstores.faiss import FAISS
from pathlib import Path
import re
import spacy
from typing import Generator

def iter_docx_data(page_id : int, page_name : str, doc : DocumentObject, nlp : spacy.Language) -> Generator[LangchainDocument,None,None]:
	containers = doc.tables.copy()

	containers = []

	for table in doc.tables:
		values = []

		for row in table.rows:
			row_values = []
			
			for cell in row.cells:
				row_values.append(cell.text)

			values.append(row_values)

		containers.append((table._element,values))
	''' Парсинг списков
	
	current_list = []
	first_element = None

	for paragraph in doc.paragraphs:
		if "List" in paragraph.style.name:
			if not current_list:
				first_element = paragraph._element

			current_list.append([paragraph.text])
		elif current_list:
			containers.append((first_element,current_list))
			current_list = []
			first_element = None
	'''
	for i in range(len(containers)):
		container_element,container_values = containers[i][0],containers[i][1]

		if len(container_values) <= 1:
			continue

		full_context = []

		for element in doc.element.body:
			if element == container_element:
				break

			if element.text:
				full_context.append(element.text)

		context_before = []

		for paragraph in full_context[-3:]:
			paragraph_sents = []
			parsed_paragraph = nlp(paragraph)
			#print('Paragraph: ', parsed_paragraph)
			for sent in parsed_paragraph.sents:
				paragraph_sents.append(sent.text)
				#print('Sent: ', sent)
			context_before.append(paragraph_sents)


		uuid = f'wiki-{page_id}-{"_".join(page_name.split())}-{i + 1}'

		header_line = " | ".join(container_values[0])

		table_text = '\n'.join([' | '.join(row) for row in container_values[1:]])

		context_text = '\n'.join('\n'.join(inner_text) for inner_text in context_before)

		full_text = f"{page_name}\n{context_text}\n{""}\n{header_line}\n{table_text}".strip()

		yield LangchainDocument(
			page_content=full_text,
			metadata={
				"uuid": uuid,
				"header": container_values[0]
			}
		)
		
def docx2faiss(files : list[str], vectorstore : FAISS, nlp : spacy.Language) -> None:
	for i in range(len(files)):
		print(f"={i + 1}/{len(files)}=: Parsing document: {files[i]}")

		filepath = Path(files[i])

		documents = []

		docx = CreateDocument(Path(files[i]))
		
		for doc in iter_docx_data(i, filepath.stem, docx, nlp):
			documents.append(doc)
			
		vectorstore.add_documents(documents)